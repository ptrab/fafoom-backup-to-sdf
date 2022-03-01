#!/usr/bin/env python
""" parse the tddft stuff of gaussian """

import sys
import argparse

# import subprocess
from pyparsing import *

ParserElement.enablePackrat()


def getinput(args):
    """parse the input"""
    parser = argparse.ArgumentParser(description="G09 to Table in docx-file converter.")
    parser.add_argument(
        "outputfile",
        metavar="G09-Output",
        help="Typically *.log or *.out, but ending doesn't matter.",
    )
    parser.add_argument(
        "--out",
        "-o",
        default="table.docx",
        help="word document for the table (default: table.docx)",
    )
    # parser.add_argument(
    #    "--fthresh",
    #    "-t",
    #    default=0.0,
    #    type=float,
    #    help="threshold for oscillator strengths",
    # )
    parser.add_argument(
        "--states",
        "-st",
        nargs="+",
        required=True,
        type=int,
        help="Specify the wanted states. Else, all states are converted.",
    )
    # s2_group = parser.add_mutually_exclusive_group(required=False)
    # s2_group.add_argument(
    #    "--s2thresh",
    #    "-c",
    #    type=float,
    #    help="threshold for <S**2> states "
    # )
    # s2_group.add_argument(
    #    "--s2mult",
    #    "-m",
    #    action="store_true",
    #    help="takes multiplicity to set <S**2>-threshold",
    # )

    return parser.parse_args(args)


def remove_last_line_from_string(string):
    """ removes the last line of a string """
    # https://stackoverflow.com/a/18683105/6155796
    result = string[: string.rfind("\n")]
    return result


def is_closed_shell(raw):
    """ is closed shell? """
    use_pyparse = False
    if use_pyparse:
        # I want to find out the Multiplicity of the system
        # which is shown in the output like that:
        # Charge =  0 Multiplicity = 2
        #
        # definition of the pyparsing words
        charge = Literal("Charge")
        mult = Literal("Multiplicity")
        eq_sign = Literal("=")
        num = Word(nums).setParseAction(tokenMap(int))
        # putting it together
        introduction = Suppress(charge + eq_sign + num + mult + eq_sign)
        find = introduction + num
        # searching for the string
        result = find.searchString(raw)[0][0]
        # setting the proper scaling factor for CI coefficients
    else:
        lines = raw.splitlines()
        for item in lines:
            if "Multiplicity" in item:
                splitted_line = item.split()
                result = int(splitted_line[-1])
                break
    if result > 1:
        # false
        scale = 1.0
    else:
        # true
        scale = 0.5
    # returning the scaling factor as well as the multiplicity
    # no need for the latter so far
    return scale, result


def num_basis_functions(raw):
    """ get the number of basis functions """
    # 952 basis functions,  1755 primitive gaussians,  1014 cartesian basis functions
    use_pyparse = False
    if use_pyparse:
        num = Word(nums).setParseAction(tokenMap(int))
        basis_functions = Literal("basis functions,")
        primitive_gaussians = Literal("primitive gaussians,")
        cartesian_basis = Literal("cartesian basis functions")
        find = num + Suppress(
            basis_functions + num + primitive_gaussians + num + cartesian_basis
        )
        result = find.searchString(raw)[0][0]
    else:
        lines = raw.splitlines()
        for item in lines:
            if "basis functions" in item and not "There are" in item:
                print(item)
                splitted_line = item.split()
                result = int(splitted_line[0])
                break
    return result


def spin_contamination(spin, s_squared):
    """ gives the contribution of the next higher contaminating state """
    # it assumes, that in spin contamination the next highest
    # lying excited state has the greatest impact
    #
    #              <s²> - S(S+1)
    # contrib(S) = -------------
    #               (S+1) (S+2)
    #
    result = (s_squared - spin * (spin + 1)) / ((spin + 1) * (spin + 2))
    return result


def parse_text(raw):
    """ parse the text """
    # closed shell:
    #  Excited State   1:      Singlet-B1    14.8877 eV   83.28 nm  f=0.0037  <S**2>=0.000
    #       5 ->  6         0.70759
    # open shell:
    #  Excited State   1:  2.005-A      0.4398 eV 2818.97 nm  f=0.0244  <S**2>=0.755
    #       222A ->223A        0.99831
    num = Word(nums)
    symmetry = Suppress(Word(alphanums + "-.?"))
    colon = Literal(":")
    state = Combine(num + Suppress(colon))
    float_ = Word(nums + ".-")
    unit = Suppress(oneOf("eV nm"))
    oscstr = Combine(Suppress("f=") + float_)
    spincont = Combine(Suppress("<S**2>=") + float_)
    arrows = oneOf("<- ->")
    excitedstate = (
        Suppress("Excited State")
        + state
        + symmetry
        + float_
        + unit
        + float_
        + unit
        + oscstr
        + spincont
    )
    excitations = (
        Combine(num + ZeroOrMore(oneOf("A B")))
        + arrows
        + Combine(num + ZeroOrMore(oneOf("A B")))
        + float_.setParseAction(tokenMap(float))
    )
    mylines = Group(excitedstate) + Group(ZeroOrMore(Group(excitations)))
    try:
        # [['3', '18.1202', '68.42', '0.0672', '0.000'], [['3', '->', '7', '0.12606'], ['4', '->', '6', '0.69577']]]
        result = mylines.searchString(raw)
        # print(result)
        # for item in result:
        #    state, excitations = item
        #    print(state)
        #    for exc in excitations:
        #        print(exc)
    # https://stackoverflow.com/a/41000491/6155796
    except ParseException as pe:
        print(pe.line)
        print(" " * (pe.col - 1) + "^")
        print(pe)
    return result


# https://github.com/python-openxml/python-docx/issues/322
def set_repeat_table_header(row):
    """ set repeat table row on every new page
    """
    from docx.oxml.shared import OxmlElement, qn
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row


def to_docx(content, scale, inputargs, basisfunctions):
    """ writes content into a ms word table """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

    # from docx.shared import Cm

    # from docx.shared import Inches

    document = Document()

    font = document.styles["Normal"].font
    font.name = "Calibri"
    font.size = Pt(8)

    # document.add_heading("My Table", 0)

    # p = document.add_paragraph("Hopefully my table will be here.")
    # p.add_run("bold").bold = True
    # p.add_run(" and some ")
    # p.add_run("italic.").italic = True

    # document.add_heading("Heading, level 1", level=1)
    # document.add_paragraph("Intense quote", style="IntenseQuote")

    # document.add_paragraph("first item in unordered list", style="ListBullet")
    # document.add_paragraph("first item in ordered list", style="ListNumber")

    # document.add_picture("monty-truth.png", width=Inches(1.25))

    # [['3', '18.1202', '68.42', '0.0672', '0.000'], [['3', '->', '7', '0.12606'], ['4', '->', '6', '0.69577']]]
    # NR_OF_STATES = len(content)

    table = document.add_table(rows=1, cols=10)
    table.allow_autofit = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Nr"
    hdr_cells[1].text = "En\neV"
    hdr_cells[2].text = "\u03BB\nnm"
    hdr_cells[3].text = "f"
    hdr_cells[4].text = "\u3008" + "S²" + "\u3009"
    hdr_cells[5].text = "Wgt\n%"
    hdr_cells[6].text = "From"
    hdr_cells[7].text = "To"
    hdr_cells[8].text = "Hole"
    hdr_cells[9].text = "Electron"
    set_repeat_table_header(table.rows[0])
    wanted = inputargs.states
    mos = []
    for item in content:
        state, excitations = item
        state.nr, state.en, state.wl, state.f, state.sc = state
        if not int(state.nr) in wanted:
            continue
        row_cells = table.add_row().cells
       # row_cells.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row_cells[0].text = state.nr
        row_cells[1].text = f"{state.en:.2f}"
        row_cells[2].text = f"{state.wl:.0f}"
        row_cells[3].text = state.f
        row_cells[4].text = state.sc
        weights = ""
        froms = ""
        tos = ""
        for entry in excitations:
            if entry[1] == "->":
                weight = 100.0 / scale * entry[3] ** 2
                if weight >= 8.0:
                    weights = weights + f"{weight:.0f}" + "\n"
                    froms = froms + entry[0] + "\n"
                    tos = tos + entry[2] + "\n"
                    mos.append(entry[0])
                    mos.append(entry[2])
        weights = remove_last_line_from_string(weights)
        froms = remove_last_line_from_string(froms)
        tos = remove_last_line_from_string(tos)
        row_cells[5].text = weights
        row_cells[6].text = froms
        row_cells[7].text = tos

        try:
            paragraph = row_cells[8].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture("hole" + state.nr + "_thumb.jpg", width=Cm(3.0))
        except:
            pass

        try:
            paragraph = row_cells[8].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture("hole" + state.nr + "_thumb.jpeg", width=Cm(3.0))
        except:
            pass

        try:
            paragraph = row_cells[8].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture("hole" + state.nr + "_thumb.png", width=Cm(3.0))
        except:
            pass

        try:
            paragraph = row_cells[9].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture("electron" + state.nr + "_thumb.jpg", width=Cm(3.0))
        except:
            pass

        try:
            paragraph = row_cells[9].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture("electron" + state.nr + "_thumb.jpeg", width=Cm(3.0))
        except:
            pass

        try:
            paragraph = row_cells[9].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture("electron" + state.nr + "_thumb.png", width=Cm(3.0))
        except:
            pass

    # document.save(inputargs.out)

    sorted_mos = sorted(list(set(mos)))
    a_mos = []
    b_mos = []
    b_mos_bf = []
    c_mos = []
    for item in sorted_mos:
        if "A" in item:
            a_mos.append(item)
        elif "B" in item:
            b_mos.append(item)
        else:
            c_mos.append(item)

    if c_mos != []:
        c_mos = [int(x) for x in c_mos]
        print(f"{min(c_mos):.0f} ... {max(c_mos):.0f}")
    else:
        a_mos = [int(x.replace("A", "")) for x in a_mos]
        b_mos = [int(x.replace("B", "")) for x in b_mos]
        b_mos_bf = [basisfunctions + x for x in b_mos]
        print(f"A: {min(a_mos):.0f} ... {max(a_mos):.0f}")
        print(f"B: {min(b_mos):.0f} ... {max(b_mos):.0f}")
        print(f"B: {min(b_mos_bf):.0f} ... {max(b_mos_bf):.0f}")

    document.add_paragraph()

    table = document.add_table(rows=0, cols=6)
    table.allow_autofit = True
    if c_mos != []:
        moRange = range(min(c_mos), max(c_mos) + 1)
        # https://stackoverflow.com/a/312464/6155796
        moRange = [moRange[i : i + 6] for i in range(0, len(moRange), 6)]
        for item in moRange:
            row_cells = table.add_row().cells
            for i in range(len(item)):
                try:
                    paragraph = row_cells[i].paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(
                        "orb{:0>6}_thumb.jpg".format(item[i]), width=Cm(2.6)
                    )
                except:
                    pass

                try:
                    paragraph = row_cells[i].paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(
                        "orb{:0>6}_thumb.jpeg".format(item[i]), width=Cm(2.6)
                    )
                except:
                    pass

                try:
                    paragraph = row_cells[i].paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(
                        "orb{:0>6}_thumb.png".format(item[i]), width=Cm(2.6)
                    )
                except:
                    pass

            row_cells = table.add_row().cells
            for i in range(len(item)):
                row_cells[i].text = str(item[i])

    else:
        moRangeA = range(min(a_mos), max(a_mos))
        moRangeA = [moRangeA[i : i + 6] for i in range(0, len(moRangeA), 6)]
        moRangeB = range(min(b_mos), max(b_mos))
        moRangeB = [moRangeB[i : i + 6] for i in range(0, len(moRangeB), 6)]
        moRangeBF = range(min(b_mos_bf), max(b_mos_bf))
        moRangeBF = [moRangeBF[i : i + 6] for i in range(0, len(moRangeBF), 6)]
        for item in moRangeA:
            row_cells = table.add_row().cells
            for i in range(len(item)):
                try:
                    paragraph = row_cells[i].paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(
                        "orb{:0>6}_thumb.jpg".format(item[i]), width=Cm(2.6)
                    )
                except:
                    pass

                try:
                    paragraph = row_cells[i].paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(
                        "orb{:0>6}_thumb.jpeg".format(item[i]), width=Cm(2.6)
                    )
                except:
                    pass

                try:
                    paragraph = row_cells[i].paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(
                        "orb{:0>6}_thumb.png".format(item[i]), width=Cm(2.6)
                    )
                except:
                    pass

            row_cells = table.add_row().cells
            for i in range(len(item)):
                row_cells[i].text = str(item[i])

        for j in range(len(moRangeB)):
            item1 = moRangeB[j]
            item2 = moRangeBF[j]
            row_cells = table.add_row().cells
            for i in range(len(item1)):
                try:
                    paragraph = row_cells[i].paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(
                        "orb{:0>6}_thumb.jpg".format(item[i]), width=Cm(2.6)
                    )
                except:
                    pass

                try:
                    paragraph = row_cells[i].paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(
                        "orb{:0>6}_thumb.jpeg".format(item[i]), width=Cm(2.6)
                    )
                except:
                    pass

                try:
                    paragraph = row_cells[i].paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(
                        "orb{:0>6}_thumb.png".format(item[i]), width=Cm(2.6)
                    )
                except:
                    pass

            row_cells = table.add_row().cells
            for i in range(len(item2)):
                row_cells[i].text = str(item2[i])

    document.save(inputargs.out)


if __name__ == "__main__":
    ARGS = getinput(sys.argv[1:])
    # FILE_CONTENT = (
    #   open(ARGS.outputfile, "r").read().replace("->", "-> ").replace("<-", "<- ")
    # )
    MY_FILE = open(ARGS.outputfile, "r").readlines()
    FILE_CONTENT = []
    for line in MY_FILE:
        if (
            "basis functions"
            or "Multiplicity"
            or "Excited State"
            or "->"
            or "<-" in line
        ):
            FILE_CONTENT.append(line.replace("->", "-> ").replace("<-", "<- "))
    FILE_CONTENT = "".join(FILE_CONTENT)
    SCALE_FACTOR, MULTIPLICITY = is_closed_shell(FILE_CONTENT)
    # formatierter_string = f"{ein_float:.2f}"
    print(
        (
            "The molecule is of Multiplicity " + str(MULTIPLICITY) + " and thus "
            "the excitations will be scaled by " + str(SCALE_FACTOR)
        )
    )
    # num_basis_functions(FILE_CONTENT)
    to_docx(
        parse_text(FILE_CONTENT), SCALE_FACTOR, ARGS, num_basis_functions(FILE_CONTENT)
    )
